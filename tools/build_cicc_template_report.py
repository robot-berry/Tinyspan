from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = Path("G:/UESTC/uav/01/PLin+SingleNet+HDMI/competition01")
OUTPUT = ROOT / "docs" / "TinySPAN_CICC_template_technical_report_20260630.docx"

BLACK = RGBColor(0, 0, 0)
FONT_CN = "宋体"
FONT_HEAD = "黑体"


def find_template() -> Path:
    matches = list(TEMPLATE_DIR.glob("CICC1000997*.docx"))
    if not matches:
        raise FileNotFoundError(f"未找到模板 DOCX: {TEMPLATE_DIR}")
    return matches[0]


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, name: str = FONT_CN, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(style, name: str, size: float, bold: bool = False) -> None:
    font = style.font
    font.name = name
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = BLACK
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    set_style_font(styles["Normal"], FONT_CN, 10.5)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        set_style_font(styles[name], FONT_HEAD, size, True)
        styles[name].paragraph_format.space_before = Pt(12)
        styles[name].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Title"], FONT_HEAD, 22, True)
    if "Subtitle" in styles:
        set_style_font(styles["Subtitle"], FONT_CN, 12, False)
    for style in styles:
        try:
            if style.font is not None:
                style.font.color.rgb = BLACK
        except Exception:
            pass


def set_cell_shading(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    set_run_font(run, FONT_CN, 9.5, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def apply_table_style(table) -> None:
    for name in ("Table Grid", "网格型", "网格表", "TableNormal"):
        try:
            table.style = name
            return
        except Exception:
            continue


def add_p(doc: Document, text: str = "", style: str | None = None, align=None, bold=False, size=None, font=FONT_CN):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    if text:
        run = para.add_run(text)
        set_run_font(run, font, size, bold)
    return para


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    set_run_font(run, FONT_HEAD, {1: 16, 2: 14, 3: 12}.get(level, 12), True)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    add_p(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, True)
        set_cell_shading(cell)
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, False, align)
            if widths:
                cells[i].width = Inches(widths[i])
    add_p(doc)


def add_numbered_items(doc: Document, items: list[str]) -> None:
    for i, text in enumerate(items, 1):
        add_p(doc, f"（{i}）{text}")


def add_picture_if_exists(doc: Document, rel_path: str, caption: str) -> None:
    path = ROOT / rel_path
    if not path.exists():
        add_p(doc, f"{caption}：图片文件未在当前工程中找到，预期路径为 {rel_path}")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Inches(5.8))
    add_p(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)


def add_cover(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    add_p(doc, "第十届", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20, font=FONT_HEAD)
    add_p(doc, "全国大学生集成电路创新创业大赛", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=22, font=FONT_HEAD)
    for _ in range(3):
        add_p(doc)
    add_p(doc, "技术报告", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=26, font=FONT_HEAD)
    for _ in range(4):
        add_p(doc)

    meta = [
        ("报告类型", "技术报告"),
        ("参赛杯赛", "端侧 AI 超分硬件加速器赛题"),
        ("作品名称", "基于 TinySPAN 的端侧视频超分硬件加速器设计与实现"),
        ("队伍编号", "待填写"),
        ("团队名称", "待填写"),
        ("工程路径", r"G:\UESTC\feitengspan1\Tinyspan"),
        ("生成日期", "2026-06-30"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (k, v) in zip(table.rows, meta):
        set_cell_text(row.cells[0], k, True)
        set_cell_shading(row.cells[0], "EAF2F8")
        set_cell_text(row.cells[1], v, False, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()


def add_toc_placeholder(doc: Document) -> None:
    add_p(doc, "目录", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, font=FONT_HEAD)
    add_p(doc, "[[TOC]]")
    doc.add_page_break()


def build_doc() -> Path:
    template = find_template()
    doc = Document(template)
    clear_body(doc)
    setup_styles(doc)

    add_cover(doc)
    add_toc_placeholder(doc)

    add_heading(doc, "快速预览简介", 1)
    add_p(doc, "作品定位")
    add_p(
        doc,
        "本作品面向端侧视频会议与图像增强场景，设计并实现一套基于 TinySPAN 的专用 AI 超分硬件加速器。系统以 Xilinx Zynq UltraScale+ xczu19eg 板卡为验证平台，按 ZC706 / XC7Z045 等效资源门线统计 PPA，目标是在 REDS 数据集相关验证链路上实现 X2、X4 超分，并输出 1280x720@30fps 的实时结果。",
    )
    add_table(
        doc,
        "表1-快速预览关键技术",
        ["技术模块", "采用方案", "作用"],
        [
            ["AI 模型", "TinySPAN C32/B4 student", "在较低 DSP、LUT、BRAM 资源下完成 X2/X4 超分"],
            ["量化策略", "W8A8 定点量化", "将训练模型转换为硬件可实现的整数计算链路"],
            ["整帧路线", "64x64 LR tile 切块", "支持 320x180 X4 与 640x360 X2 输入拼接为 720p 输出"],
            ["输入输出", "PS DDR + PL tile pipeline", "避免低速全帧 JTAG 读回影响实时吞吐判断"],
            ["验收方法", "board-vs-fixed 字节级比较", "保证真实板上输出与同 checkpoint 定点参考一致"],
        ],
    )
    add_table(
        doc,
        "表2-快速预览核心指标",
        ["指标", "X4", "X2"],
        [
            ["LR 输入", "320x180", "640x360"],
            ["SR 输出", "1280x720", "1280x720"],
            ["tile 组织", "64x64 LR，15 tile", "64x64 LR，60 tile"],
            ["实测吞吐", "30.409639fps @155MHz", "32.860482fps @187.512MHz"],
            ["正确性", "0 / 2764800 mismatch，max diff 0", "0 / 2764800 mismatch，max diff 0"],
            ["功耗", "Total On-Chip 3.969W", "Total On-Chip 4.053W"],
        ],
    )
    add_p(doc, "核心亮点")
    add_p(
        doc,
        "本方案的核心亮点在于以较小模型和明确的定点参考闭合真实上板验证：X2 与 X4 两个倍率均完成板上全帧切块超分、DDR 内字节级对比和 30fps 以上吞吐验收；资源占用远低于 ZC706 等效门线，便于后续按 AXI DMA、VDMA 或 DataMover 升级输入输出链路。",
    )

    add_heading(doc, "摘要", 1)
    add_p(
        doc,
        "面向端侧场景的 AI 超分任务需要在清晰度、实时性和功耗之间取得平衡。传统 GPU 或 CPU 方案不适合直接部署到资源受限 FPGA，而过大的超分网络又会带来 DSP、BRAM 和时序压力。本文设计并验证了一套 TinySPAN 专用超分硬件加速器，以 TinySPAN C32/B4 student 作为提交模型，通过 W8A8 量化、定点参考、RTL 参数导出和真实板上比较构成完整交付链路。",
    )
    add_p(
        doc,
        "系统采用整帧切块上板路线：LR 帧进入 PS DDR 后，PL 按固定 64x64 LR tile 取数并执行超分，对边界 tile 做 zero padding 和有效区域裁剪，最终写回完整 720p SR 帧。该路线避免一次性展开大尺寸输入造成资源和时序不可控，同时保留整图/视频帧实时处理能力。",
    )
    add_p(
        doc,
        "当前提交以已有 Gate H 硬件闭合结果作为比赛基线。X4 在 320x180 LR 输入下输出 1280x720，实测 30.409639fps；X2 在 640x360 LR 输入下输出 1280x720，实测 32.860482fps。两者均达到 board-vs-fixed 0 mismatch，满足本阶段以正确性、实时性和低资源为优先的提交要求。",
    )

    add_heading(doc, "一、赛题目标与需求分析", 1)
    add_heading(doc, "1.1 赛题目标", 2)
    add_p(
        doc,
        "赛题要求面向端侧场景设计专用 AI 超分硬件加速器，实现清晰度、实时性和功耗之间的平衡。模型结构、量化和稀疏化方法不做限制，图像输入格式可采用 RGB 或 YUV；方案需要适配会议画面中的人像、文档等核心特征，并在 Xilinx ZC706 或资源相当 FPGA 上对 REDS 数据集实现 X2、X4 超分验证。",
    )
    add_heading(doc, "1.2 功能需求", 2)
    add_numbered_items(
        doc,
        [
            "提供 AI 模型结构、训练、量化说明文档及源代码。",
            "提供模型到硬件加速器指令或参数的转换工具。",
            "提供硬件加速器详细设计文档和可综合 RTL 源代码。",
            "提供 Vivado 仿真、综合、实现和真实上板验证证据。",
            "提供图像质量、实时吞吐、资源面积和功耗分析。",
        ],
    )
    add_heading(doc, "1.3 性能与资源需求", 2)
    add_p(
        doc,
        "本工程将实时输出目标定义为 1280x720@30fps。为了贴合 ZC706 等效资源门线，资源统计按 LUT 218600、Register 437200、DSP 900、BRAM Tile 545 作为归一化参考，同时记录 xczu19eg 上的 timing、频率和功耗结果。",
    )
    add_table(
        doc,
        "表3-赛题要求对应关系",
        ["赛题要求", "本提交对应内容", "状态"],
        [
            ["AI 模型、训练、量化说明", "docs/model_design.md、docs/training_quantization.md、train/、configs/", "已提供"],
            ["模型到硬件转换工具", "tools/model_to_hardware/ 下的 quant plan、integer reference、RTL export 工具", "已提供"],
            ["硬件设计文档", "docs/hardware_design.md、docs/full_frame_tiling_route.md、docs/verification_plan.md", "已提供"],
            ["硬件源代码", "rtl/tinyspan_core/、rtl/board_wrapper/、scripts/vivado/", "已提供"],
            ["X2/X4 实时验证", "X2/X4 Gate H 上板 manifest 与 DDR 内 compare 结果", "已闭合"],
            ["PPA 分析", "docs/ppa_analysis.md 与 Vivado utilization/timing/power 报告", "已提供"],
        ],
    )

    add_heading(doc, "二、系统总体设计", 1)
    add_heading(doc, "2.1 总体架构", 2)
    add_p(
        doc,
        "系统由训练与量化链路、模型到硬件参数转换链路、PL 端 TinySPAN tile pipeline、PS DDR 输入输出链路和验证脚本组成。软件侧冻结 checkpoint 并生成定点参考；硬件侧读取导出的权重、scale 和 requant 参数，按 tile 执行超分；上板侧以 A53 在 DDR 内进行输出与 reference 的逐字节比较。",
    )
    add_table(
        doc,
        "表4-系统模块划分",
        ["模块", "输入", "输出", "功能定位"],
        [
            ["训练与冻结", "REDS 数据集、student 配置", "frozen checkpoint", "形成可提交模型参数"],
            ["量化与导出", "checkpoint、calibration", "W8A8 quant plan、RTL memory", "连接 PyTorch 与 RTL"],
            ["TinySPAN pipeline", "64x64 LR tile、权重、scale", "SR tile", "完成定点超分计算"],
            ["帧切块控制", "LR frame 地址、尺寸、倍率", "tile read/write schedule", "完成整帧拆分和边界裁剪"],
            ["板上验收", "DDR 输入、fixed reference", "mismatch、fps、manifest", "确认真实硬件输出正确性和吞吐"],
        ],
    )
    add_heading(doc, "2.2 整帧切块路线", 2)
    add_p(
        doc,
        "本工程不要求硬件一次性接收完整 720p 大帧进入卷积阵列，而是按 tile 流水处理。X4 场景中，320x180 LR 帧切为 5x3 个 64x64 LR tile；X2 场景中，640x360 LR 帧切为 10x6 个 64x64 LR tile。对右边界和下边界不足 64 的 tile 做 zero padding，写回时只保留有效 SR 区域。",
    )
    add_heading(doc, "2.3 输入输出路线", 2)
    add_p(
        doc,
        "当前严格验收采用 PS DDR 输入、PL 运行、A53 DDR 内比较的方式，避免慢速 JTAG 全帧读回成为吞吐瓶颈。后续工程化视频输入输出可直接调用板卡 PS DDR controller、AXI HP/HPC 端口和 Xilinx 标准 AXI DMA、VDMA 或 DataMover IP，不自研 DDR controller/PHY。",
    )

    add_heading(doc, "三、AI 模型、训练与量化", 1)
    add_heading(doc, "3.1 TinySPAN 模型结构", 2)
    add_numbered_items(
        doc,
        [
            "模型结构：TinySPAN C32/B4，32 个特征通道，4 个 TinySPAN block。",
            "输入格式：RGB888 LR frame 或 LR tile。",
            "输出格式：RGB888 SR frame。",
            "量化格式：W8A8，权重 INT8、激活 INT8，bias/requant 参数由 quant plan 固化。",
            "X4 硬件安全基线：c32b4_30fps_frozen_20260613。",
            "X2 提交 checkpoint：runs/tinyspan_frozen_candidates/x2_quality_after_x4_20260625/student_final.pt。",
        ],
    )
    add_heading(doc, "3.2 数据集与训练", 2)
    add_p(
        doc,
        "训练使用 REDS 数据集，仓库不包含 REDS 原始图片，训练脚本通过参数指定本地或云端数据路径。当前提交不继续等待 X4/X2 画质提升训练支线，而是使用已经完成真实上板闭合的 TinySPAN Gate H 方案作为比赛基线。",
    )
    add_table(
        doc,
        "表5-X2 软件质量评估",
        ["指标", "TinySPAN X2", "Bicubic X2", "增益"],
        [
            ["PSNR mean", "31.121459919 dB", "30.853986135 dB", "+0.267473784 dB"],
            ["SSIM mean", "0.905514798", "0.899959173", "+0.005555626"],
            ["MAE/255 mean", "0.017395369", "0.017778457", "更低"],
        ],
    )
    add_p(
        doc,
        "X4 当前提交基线强调真实板上正确性、实时性和低资源。已尝试的 X4 画质提升候选 QUALITY_X4_HRHEAVY_P256_20260626 在 full REDS val 上约为 26.384690522dB，未达到 28dB 提升目标，因此未替换当前 X4 Gate H 提交基线。",
    )
    add_heading(doc, "3.3 量化与硬件参数转换", 2)
    add_p(doc, "模型到硬件转换链路如下：")
    add_numbered_items(
        doc,
        [
            "冻结 PyTorch student checkpoint。",
            "融合 TinySPAN checkpoint 并生成 manifest。",
            "执行 calibration 并确定 activation scale。",
            "生成 W8A8 quant plan。",
            "生成软件定点 reference。",
            "导出 RTL memory、权重和 requant 参数。",
            "运行 Vivado 仿真、实现和 bitstream 生成。",
            "执行 board-vs-fixed 上板验收。",
        ],
    )
    add_table(
        doc,
        "表6-关键模型与量化文件校验",
        ["对象", "SHA256"],
        [
            ["X4 checkpoint", "6A3AA4FE17CDF1027483F95BE8A99A5805BCDD61CC821074603DE65BF333D938"],
            ["X4 quant plan", "EB6EEDDDE9360F61E6FC30141B2A1E6539E519CB226AC18B8C219B9E40092C9D"],
            ["X2 checkpoint", "B06E66FA8FEA066F111B94CF5629919BEA05D5465F913B36851CBA92BED4A9EB"],
            ["X2 quant plan", "8BB154CC524B5CA00A1A3D81F0E343CF0A0BA1CF8E08AB5867656B0C53D37C2F"],
        ],
    )

    add_heading(doc, "四、硬件加速器设计", 1)
    add_heading(doc, "4.1 硬件模块划分", 2)
    add_table(
        doc,
        "表7-硬件加速器模块",
        ["模块", "功能"],
        [
            ["PS/DDR 入口", "通过板卡 PS DDR controller IP 和 AXI HP/HPC 端口访问输入/输出帧，不自研 DDR controller/PHY"],
            ["TinySPAN tile pipeline", "按 64x64 LR tile 读取 RGB，执行 W8A8 TinySPAN 卷积、激活、重排与 SR tile 输出"],
            ["动态裁剪与写回", "对边界 tile 做 zero padding 输入和 valid region crop 输出，写回完整 720p SR frame"],
            ["控制与计数", "记录 frame cycles、吞吐 fps、tile 调度和完成状态"],
        ],
    )
    add_heading(doc, "4.2 时序与资源约束", 2)
    add_p(
        doc,
        "X4 Gate H 在 155MHz 下 WNS 为 +0.020ns，X2 Gate H 在 187.512MHz 下 WNS 为 +0.002ns，两者均达到 timing pass。设计没有使用 URAM，DSP、BRAM 和 LUT 占用均低于 ZC706 等效资源门线。",
    )
    add_heading(doc, "4.3 与大尺寸输入的关系", 2)
    add_p(
        doc,
        "设计不依赖一次性展开 160x90、320x180 或 720p 大尺寸卷积窗口。整帧实时性能来自 tile 级流水、DDR 帧缓存和有效区域拼接，因此只要 tile 调度和输入输出带宽满足要求，就可以面向完整图像或视频帧实现实时超分。",
    )

    add_heading(doc, "五、系统实现方案", 1)
    add_heading(doc, "5.1 工程路径与主要目录", 2)
    add_table(
        doc,
        "表8-交付文件索引",
        ["类别", "路径"],
        [
            ["工作流", "WORKFLOW.md"],
            ["模型结构", "docs/model_design.md"],
            ["训练与量化", "docs/training_quantization.md"],
            ["硬件设计", "docs/hardware_design.md"],
            ["验证方案", "docs/verification_plan.md"],
            ["PPA 分析", "docs/ppa_analysis.md"],
            ["模型源码", "train/"],
            ["转换工具", "tools/model_to_hardware/"],
            ["RTL 源码", "rtl/tinyspan_core/、rtl/board_wrapper/"],
            ["Vivado 脚本", "scripts/vivado/"],
            ["X4 上板证据", "artifacts/.../gate_h_board_x4_320x180_f150_tiledref_tile64_fifo_f155_20260625/"],
            ["X2 上板证据", "artifacts/.../gate_h_board_x2_640x360_f188_div8_tile64_rgbpipe_20260626/"],
        ],
        widths=[1.7, 5.5],
    )
    add_heading(doc, "5.2 上板执行流程", 2)
    add_numbered_items(
        doc,
        [
            "准备 LR 输入帧、fixed reference 和 quant plan。",
            "通过 PS DDR 写入输入帧和控制参数。",
            "启动 PL TinySPAN tile pipeline。",
            "等待硬件完成并读取 frame cycles、done 和状态寄存器。",
            "由 A53 在 DDR 内逐字节比较 board output 与 fixed reference。",
            "生成 manifest、comparison preview、diff heatmap 和 PPA 记录。",
        ],
    )

    add_heading(doc, "六、测试方案与结果分析", 1)
    add_heading(doc, "6.1 验证层级", 2)
    add_table(
        doc,
        "表9-验证方案",
        ["层级", "验证内容", "证据"],
        [
            ["软件训练质量", "REDS val_sharp PSNR/SSIM/MAE", "X2 quality candidate manifest"],
            ["量化一致性", "PyTorch student、定点 reference、integer reference 比较", "tools/image_validation/ 输出"],
            ["RTL/仿真", "tile wrapper、cropper、writer shell、PS/DDR wrapper", "sim/reports/"],
            ["Vivado 实现", "bitstream、timing、utilization、power", "Vivado reports 与 Gate H manifest"],
            ["真实上板", "A53 in-DDR full-frame compare、frame cycles、fps", "X2/X4 Gate H manifest"],
        ],
    )
    add_heading(doc, "6.2 X4 上板结果", 2)
    add_table(
        doc,
        "表10-X4 Gate H 实测结果",
        ["项目", "结果"],
        [
            ["输入/输出", "320x180 LR -> 1280x720 SR"],
            ["tile", "64x64 LR，5x3=15 tile"],
            ["频率", "155MHz"],
            ["吞吐", "30.409639424fps"],
            ["正确性", "board-vs-fixed 0 / 2764800 mismatch，max diff 0"],
            ["WNS/WHS", "+0.020ns / +0.007ns"],
            ["功耗", "Total On-Chip 3.969W，Dynamic 2.755W"],
        ],
    )
    add_heading(doc, "6.3 X2 上板结果", 2)
    add_table(
        doc,
        "表11-X2 Gate H 实测结果",
        ["项目", "结果"],
        [
            ["输入/输出", "640x360 LR -> 1280x720 SR"],
            ["tile", "64x64 LR，10x6=60 tile"],
            ["频率", "187.512MHz"],
            ["吞吐", "32.860482270fps"],
            ["正确性", "board-vs-fixed 0 / 2764800 mismatch，max diff 0"],
            ["WNS/WHS", "+0.002ns / +0.014ns"],
            ["功耗", "Total On-Chip 4.053W，Dynamic 2.839W"],
        ],
    )
    add_heading(doc, "6.4 PPA 结果", 2)
    add_table(
        doc,
        "表12-PPA 指标",
        ["指标", "X4 Gate H", "X4 占 ZC706 门线", "X2 Gate H", "X2 占 ZC706 门线"],
        [
            ["CLB LUT", "6353", "2.91%", "6647", "3.04%"],
            ["CLB Register", "4647", "1.06%", "5031", "1.15%"],
            ["DSP", "81", "9.00%", "100", "11.11%"],
            ["BRAM Tile", "27", "4.95%", "27", "4.95%"],
            ["URAM", "0", "0", "0", "0"],
            ["WNS", "+0.020ns", "PASS", "+0.002ns", "PASS"],
            ["WHS", "+0.007ns", "PASS", "+0.014ns", "PASS"],
            ["PL frequency", "155MHz", "-", "187.512MHz", "-"],
            ["Total On-Chip Power", "3.969W", "-", "4.053W", "-"],
            ["Dynamic Power", "2.755W", "-", "2.839W", "-"],
        ],
    )
    add_heading(doc, "6.5 可查看图像材料", 2)
    add_picture_if_exists(
        doc,
        "artifacts/20260618_x4_tinyspan_c32b4_baseline_30fps_safe/full_frame_tiled_reference_x4_320x180_tile64_fifo_f155_20260625/comparison_preview.png",
        "图1-X4 software fixed-point 与可视化预览",
    )
    add_picture_if_exists(
        doc,
        "artifacts/20260618_x4_tinyspan_c32b4_baseline_30fps_safe/gate_h_board_x2_640x360_f188_div8_tile64_rgbpipe_20260626/tinyspan_board_software_preview.png",
        "图2-X2 board/software 对比预览",
    )
    add_p(
        doc,
        "说明：最终硬件证据是 A53 in-DDR full-frame byte compare。部分可视化图是在 0 mismatch 证明后复制 fixed reference 作为等价视图，以避免慢速全帧 JTAG 读回影响实时性判断。",
    )

    add_heading(doc, "七、作品核心亮点", 1)
    add_numbered_items(
        doc,
        [
            "X2 与 X4 均完成真实板上 720p30 级别闭合，且 board-vs-fixed 字节级完全一致。",
            "TinySPAN C32/B4 W8A8 结构资源占用低，DSP、BRAM、LUT 均远低于 ZC706 等效门线。",
            "整帧切块路线避免大尺寸卷积阵列一次性展开，兼顾资源、时序和完整帧处理能力。",
            "模型、量化、定点参考、RTL 导出、Vivado 实现和上板验收证据可追溯。",
            "输入输出路线可平滑升级到 AXI DMA、VDMA 或 DataMover，以提升工程化视频流吞吐。",
        ],
    )

    add_heading(doc, "八、风险分析与后续优化", 1)
    add_heading(doc, "8.1 当前边界", 2)
    add_numbered_items(
        doc,
        [
            "当前提交采用已有 X2/X4 Gate H 硬件闭合方案，不继续等待画质提升训练支线。",
            "X4 当前基线满足真实板上 >=30fps 和 0 mismatch，但 full REDS val 画质提升有限。",
            "当前严格验收没有把板上 SD 卡直接读图作为闭合证据；正式证据为 PS DDR 输入、PL 运行、A53 DDR 内比较。",
        ],
    )
    add_heading(doc, "8.2 后续优化方向", 2)
    add_numbered_items(
        doc,
        [
            "采用更强训练服务器继续提升 X4 / X2 画质，但新模型必须重新完成冻结、量化、RTL/export、bitstream、0 mismatch 和 >=30fps 验收。",
            "将输入输出工程化升级为 AXI DMA、VDMA 或 DataMover，减少 PS 参与和软件搬运开销。",
            "针对会议人像和文档边缘补充蒸馏、感知损失或轻量重参数化模块，在 PPA 门线内提升 PSNR/SSIM 和主观质量。",
        ],
    )

    add_heading(doc, "九、总结", 1)
    add_p(
        doc,
        "本文完成了 TinySPAN 端侧视频超分硬件加速器的模型、量化、硬件、验证和 PPA 说明。当前提交以真实板上闭合为核心依据，X2 和 X4 均达到 720p30 级别输出，输出与同一 frozen checkpoint 和 quant plan 生成的软件定点参考逐字节一致，且资源占用远低于 ZC706 等效门线。该方案可作为赛题初赛提交基线，并为后续画质提升和工程化视频输入输出升级提供稳定基础。",
    )

    add_heading(doc, "十、AI 工具使用声明", 1)
    add_p(
        doc,
        "本文档编写、格式整理和部分表述润色过程中使用了 AI 辅助工具；核心工程数据、上板结果、资源统计、训练指标和文件路径均来自当前 TinySPAN 工程记录与本地验证产物。AI 工具未替代真实板上测试、Vivado 实现或模型训练验证。",
    )

    add_heading(doc, "参考文献", 1)
    add_numbered_items(
        doc,
        [
            "S. Nah et al., NTIRE 2019 Challenge on Video Deblurring and Super-Resolution: Dataset and Study.",
            "REDS Dataset: https://seungjunnah.github.io/Datasets/reds.html",
            "CVPR NTIRE Video Super-Resolution 相关论文与开源实现。",
            "Xilinx Zynq UltraScale+ MPSoC 与 Vivado Design Suite 官方文档。",
        ],
    )

    add_heading(doc, "附录：关键工程参数", 1)
    add_table(
        doc,
        "表13-关键工程参数",
        ["参数", "取值"],
        [
            ["工程根目录", r"G:\UESTC\feitengspan1\Tinyspan"],
            ["提交路线", "TinySPAN C32/B4 W8A8 Gate H 硬件闭合方案"],
            ["目标输出", "1280x720@30fps"],
            ["X4 输入", "320x180 LR"],
            ["X2 输入", "640x360 LR"],
            ["tile 大小", "64x64 LR"],
            ["ZC706 等效门线", "LUT 218600 / Register 437200 / DSP 900 / BRAM Tile 545"],
            ["X4 manifest", "artifacts/.../gate_h_board_x4_320x180_f150_tiledref_tile64_fifo_f155_20260625/manifest.json"],
            ["X2 manifest", "artifacts/.../gate_h_board_x2_640x360_f188_div8_tile64_rgbpipe_20260626/manifest.json"],
        ],
        widths=[1.8, 5.4],
    )

    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.font.color.rgb = BLACK
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = BLACK

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = build_doc()
    print(out)
